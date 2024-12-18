from flask import Flask, render_template, request, redirect, url_for, send_file
from arxiv_search_lib import search_arxiv_papers, evaluate_papers
from database import initialize_db, save_query, get_cached_results, get_recent_queries

import xml.etree.ElementTree as ET
import io
import os
import requests
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO

import noname_library
from noname_library import noname_function


app = Flask(__name__)

# Initialize the database
initialize_db()

@app.route("/", methods=["GET", "POST"])
def search():
    if request.method == "POST":
        query = request.form.get("query")
        max_results = int(request.form.get("max_results", 10))
        return redirect(url_for("results", query=query, max_results=max_results))

    # Fetch recent queries
    recent_queries = get_recent_queries(limit=100)
    return render_template("search.html", recent_queries=recent_queries)

@app.route("/results", methods=["GET"])
def results():
    query = request.args.get("query")
    max_results = int(request.args.get("max_results", 10))

    # Check for cached results
    cached_results = get_cached_results(query, max_results)
    if cached_results:
        results = cached_results
    else:
        # Fetch and evaluate papers
        papers = search_arxiv_papers(query, max_results=max_results)
        results = evaluate_papers(papers, query)

        # Save the query and results to the database
        save_query(query, max_results, results)

    return render_template("results.html", query=query, results=results)

@app.route("/download_summary", methods=["GET"])
def download_summary():
    query = request.args.get("query")
    max_results = int(request.args.get("max_results", 10))
    
    # Fetch cached results
    cached_results = get_cached_results(query, max_results)
    if not cached_results:
        return "No cached results found for this query.", 404

    # Extract top 10 results
    top_results = cached_results[:10]

    # Create a Word document
    document = Document()
    document.add_heading(f"Top 10 Papers Summary for '{query}'", level=1)

    for idx, (pdf_url, score) in enumerate(top_results, start=1):
        try:
            # Fetch PDF
            response = requests.get(pdf_url, stream=True)
            response.raise_for_status()

            # Load PDF
            pdf_reader = PdfReader(BytesIO(response.content))
            document.add_heading(f"Paper {idx} (Score: {score})", level=2)

            # Extract 3 pages (if available)
            for page_num in range(min(3, len(pdf_reader.pages))):
                page_text = pdf_reader.pages[page_num].extract_text()
                document.add_paragraph(page_text)

        except Exception as e:
            document.add_paragraph(f"Error processing Paper {idx}: {e}")

    # Save document to a BytesIO object
    output = BytesIO()
    document.save(output)
    output.seek(0)

    # Serve the Word document
    filename = f"{query.replace(' ', '_')}_summary.docx"
    return send_file(output, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/export_history", methods=["GET"])
def export_history():
    """
    Export search history as an XML document.
    """
    # Get all queries and results from the database
    queries = get_recent_queries(limit=100)

    # Create the root element
    history = ET.Element("history")

    for query, max_results, timestamp in queries:
        query_element = ET.SubElement(history, "query")

        # Add query text
        ET.SubElement(query_element, "query_text").text = query
        ET.SubElement(query_element, "max_results").text = str(max_results)
        
        # Add results
        results = ET.SubElement(query_element, "results")

        # Get the actual results from the database
        cached_results = get_cached_results(query, max_results)
        if cached_results:
            for link, score in cached_results:
                result = ET.SubElement(results, "result")
                ET.SubElement(result, "link").text = link
                ET.SubElement(result, "score").text = str(score)

    # Convert the tree to a string
    tree = ET.ElementTree(history)
    xml_bytes = io.BytesIO()
    tree.write(xml_bytes, encoding="utf-8", xml_declaration=True)
    xml_bytes.seek(0)

    # Return the XML as a downloadable file
    return send_file(
        xml_bytes,
        as_attachment=True,
        download_name="search_history.xml",
        mimetype="application/xml"
    )


@app.route("/import_history", methods=["POST"])
def import_history():
    """
    Import search history from an XML document.
    """
    file = request.files.get("file")
    if not file:
        return "No file provided", 400

    # Parse the XML file
    try:
        tree = ET.parse(file)
        root = tree.getroot()

        # Process each query in the XML
        for query_element in root.findall("query"):
            query_text = query_element.find("query_text").text
            max_results = int(query_element.find("max_results").text)
            
            # Save query to database (to keep track of it)
            save_query(query_text, max_results, [])

            # Process results and save to database
            results = query_element.find("results")
            for result_element in results.findall("result"):
                link = result_element.find("link").text
                score = float(result_element.find("score").text)
                # Save each result
                save_query(query_text, max_results, [(link, score)])

        return "History imported successfully.", 200
    except ET.ParseError:
        return "Invalid XML file", 400

if __name__ == "__main__":
    app.run(debug=True)
